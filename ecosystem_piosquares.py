import matplotlib.pyplot as plt
from matplotlib.offsetbox import OffsetImage, AnnotationBbox
from matplotlib.gridspec import GridSpec
import matplotlib.image as mpimg
import numpy as np
import pandas as pd
import xlrd, csv, requests, shutil
from docx import Document
from docx.shared import Inches
from PIL import Image
from docx.shared import Cm
from docx.enum.text import WD_BREAK

# Step 1 - Extract variables from Json file and create dictionaries
# Convert from Excel to CSV, remove unneeded columns, only leave 1 industry
def csv_from_excel():
    wb = xlrd.open_workbook('pioneers.xls')
    sh = wb.sheet_by_index(0)
    csv_file = open('csv_file.csv', 'w')
    wr = csv.writer(csv_file, quoting=csv.QUOTE_ALL)
    for rownum in range(sh.nrows):
        wr.writerow(sh.row_values(rownum))
    csv_file.close()
    
def remove_cols():
    f=pd.read_csv("csv_file.csv")
    keep_col = ['Company Name', 'Domain', 'Description', 'Logo', 'Your industry']
    new_f = f[keep_col]
    new_f = new_f.dropna()
    new_f.to_csv("new_csv.csv", index=False)
    
csv_from_excel()
remove_cols()
file = pd.read_csv('new_csv.csv')

def coma_remove(item):
    item = item.split(',', 1)[0].replace(',', '')
    return str(item)

file['Your industry'] = file['Your industry'].apply(coma_remove)
print(file.head(20))

global industries
industries = []
for i in file['Your industry']:
    if len(industries) == 4:
        break
    if i not in industries:
        industries.append(i)
print(industries)


#Step 2 - Plot the graphs, iterate dictionaries, get logos, input into table
class graph():
    def boxes(self):        
        fig = plt.figure(constrained_layout=True, dpi=200)
        gs = GridSpec(2, 2, figure=fig)
        
        ax1 = fig.add_subplot(gs[0, 0], title = industries[0])
        industry = 0
        self.image_placer(ax1, industry)
            
        ax2 = fig.add_subplot(gs[0, 1], title = industries[1])
        industry = 1
        self.image_placer(ax2, industry)

        ax3 = fig.add_subplot(gs[1, 0], title = industries[2])
        industry = 2
        self.image_placer(ax3, industry)
        
        ax4 = fig.add_subplot(gs[1, 1], title = industries[3])
        industry = 3
        self.image_placer(ax4, industry)
        
        #fig.suptitle("Ecosystem Map")
        for i, ax in enumerate(fig.axes):
            ax.tick_params(labelbottom=False, labelleft=False, bottom = False, left = False)
        plt.savefig('ecosystem_map.png',bbox_inches='tight')
        plt.show()

    def image_placer(self, _, industry):
        count = 1
        new = file.loc[file['Your industry'] == industries[industry]]
        basewidth = 35
        for i in new['Logo']:
            png = i.split('/', 1)[-1].replace('/', '')
            r = requests.get(i, stream  = True)
            with open(png, 'wb') as out_file:
                shutil.copyfileobj(r.raw, out_file) 
            del r
            img = Image.open(png)
            wpercent = (basewidth/float(img.size[0]))
            hsize = int((float(img.size[1])*float(wpercent)))
            img = img.resize((basewidth,hsize), Image.ANTIALIAS)
            img.save(png) 
            arr_AS = mpimg.imread(png)
            imagebox = OffsetImage(arr_AS)
            if count == 1:
                x, y  = 0.2, 0.2
            if count > 1 and count < 4:
                x += 0.3
            if count  == 4:
                x = 0.2
                y = 0.5
            if count > 4 and count < 7:
                x += 0.3
            if count  == 7:
                x = 0.2
                y = 0.8
            if count > 8 and count < 10:
                x += 0.3
            if count == 10:
                x, y  = 0.2, 0.2
                count = 1
                break
            AS = AnnotationBbox(imagebox, (x, y))
            count += 1
            _.add_artist(AS)

g = graph()
g.boxes()

# Step 3 - Create Text and input all files and text into Word document 
document = Document()
p = document.add_paragraph()
p.alignment = 1
r = p.add_run()
r.add_text('Ecosystem Report for {0}, {1}, {2} and {3}'.format(*industries))
p_empty = document.add_paragraph()
p1 = document.add_paragraph()
p1.alignment = 1
r = p1.add_run()
r.add_picture('ecosystem_map.png') 
p_empty = document.add_paragraph()
p3 = document.add_paragraph()
r = p3.add_run()

def set_column_width(column, width):
    column.width = width
    for cell in column.cells:
        cell.width = width

def table_maker(_):
    p = document.add_paragraph()
    r = p.add_run()
    r.bold = True
    r.add_text(industries[_])
    table = document.add_table(1, 3)
    
    heading_cells = table.rows[0].cells
    heading_cells[0].text = 'Company Name  '
    heading_cells[1].text = 'Domain  '
    heading_cells[2].text = 'Description'
    heading_cells[2].width = Inches(1)
    new = file.loc[file['Your industry'] == industries[_]]
    keep = ['Company Name', 'Domain', 'Description']
    for name, domain, descr in new[keep].itertuples(index=False): 
        cells = table.add_row().cells
        if len(name) > 13:
            name = name[0:13] + '.. '
            cells[0].text = name
        if len(name) <= 13:
            cells[0].text = name
        cells[1].text = domain + ' '
        if len(descr) > 40:
            descr = descr[0:41] + '...'
            cells[2].text = descr
        if len(descr) <= 40:
            cells[2].text = descr   
    set_column_width(table.columns[0], Cm(8))
    set_column_width(table.columns[1], Cm(8))
    set_column_width(table.columns[2], Cm(40))
    document.add_paragraph()

for i in range(0, 3):
    table_maker(i)
    
document.save('demo.docx')