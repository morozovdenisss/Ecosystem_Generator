#!/usr/bin/env python
import matplotlib.pyplot as plt
from matplotlib.offsetbox import OffsetImage, AnnotationBbox
from matplotlib.gridspec import GridSpec
import matplotlib.image as mpimg
import numpy as np
import pandas as pd
import geopandas as gpd
import xlrd, csv, requests, shutil, json
from docx import Document
from docx.shared import Inches
from PIL import Image
from docx.shared import Pt
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from wordcloud import WordCloud, STOPWORDS
from bokeh.plotting import figure
from bokeh.models import GeoJSONDataSource, LinearColorMapper, ColorBar
from bokeh.palettes import brewer
from bokeh.io import export_png
import seaborn as sns

# Step 1
# Convert from Excel to CSV, remove unneeded columns & adapt the dataset.
def csv_from_excel():
    wb = xlrd.open_workbook('pioneers.xls')
    sh = wb.sheet_by_index(0)
    csv_file = open('csv_files/csv_file.csv', 'w')
    wr = csv.writer(csv_file, quoting=csv.QUOTE_ALL)
    for rownum in range(sh.nrows):
        wr.writerow(sh.row_values(rownum))
    csv_file.close()
    
def remove_cols():
    f=pd.read_csv("csv_files/csv_file.csv")
    keep_col = ['Company Name', 'Domain', 'Description', 'Logo', 'Your industry', 'What is the current stage of your startup?', 'Total funding received in €', 'Customer focus', 'Product focus', 'Country of incorporation / registration']
    new_f = f[keep_col]
    new_f = new_f.dropna()
    new_f.to_csv("csv_files/new_csv.csv", index=True)
    
def coma_remove(item):
    if ',' in item:
        count = 0
        iter_list = item.split(',', 1)
        size = []
        for i in range(len(iter_list) - 1):
            item = item.split(',', 1)[count].replace(',', '')
            size.append(file[file['Your industry'] == i])
            count +=  1
        item = item.split(',', 1)[size.index(max(size))].replace(',', '')
    return str(item)

# Step 2 
# All of the graphs creation & data scraping is happening here. 
class graph():
    def boxes(self):        
        fig = plt.figure(constrained_layout=True, dpi=300)
        gs = GridSpec(2, 2, figure=fig)
        ax1 = fig.add_subplot(gs[0, 0])
        industry = 0
        self.image_placer(ax1, industry)
        ax2 = fig.add_subplot(gs[0, 1])
        industry = 1
        self.image_placer(ax2, industry)
        ax3 = fig.add_subplot(gs[1, 0])
        industry = 2
        self.image_placer(ax3, industry)
        ax4 = fig.add_subplot(gs[1, 1])
        industry = 3
        self.image_placer(ax4, industry)
        #fig.suptitle("Ecosystem Map")
        for i, ax in enumerate(fig.axes):
            ax.tick_params(labelbottom=False, labelleft=False, bottom = False, left = False)
        plt.savefig('images/ecosystem_map.png',bbox_inches='tight')
    
    def image_placer(self, _, industry):
        _.set_title(industries[industry], fontsize = 8)
        count = 1
        new = file.loc[file['Your industry'] == industries[industry]]
        size = 30, 30
        for i in new['Logo']:
            png = i.split('/', 1)[-1].replace('/', '')
            r = requests.get(i, stream  = True)
            with open('logos/'+png, 'wb') as out_file:
                shutil.copyfileobj(r.raw, out_file) 
            del r
            img = Image.open('logos/'+png)
            img.thumbnail(size, Image.ANTIALIAS)
            '''img.save('logos/'+png) 
            arr_AS = mpimg.imread('logos/'+png)'''
            imagebox = OffsetImage(img)
            if count == 1:
                x, y  = 0.2, 0.2
            elif count > 1 and count < 4:
                x += 0.3
            elif count  == 4:
                x = 0.2
                y = 0.5
            elif count > 4 and count < 7:
                x += 0.3
            elif count  == 7:
                x = 0.2
                y = 0.8
            elif count > 8 and count < 10:
                x += 0.3
            elif count == 10:
                x, y  = 0.2, 0.2
                count = 1
                break
            AS = AnnotationBbox(imagebox, (x, y), frameon=False, pad=0)
            count += 1
            _.add_artist(AS)
            
    def industry_graph(self):
        plt.clf()
        plt.close()
        labels = []
        global overal_number
        global numbers_industry
        overal_number = 0
        for i in industries:
            if '&' in i:
                labels.append(i.replace(' & ', '\n').replace(' ', '\n'))
            else:
                labels.append(i.replace(' ', '\n'))
        real_labels = industries
        numbers_industry = []
        x = np.arange(len(labels))
        for i in real_labels:
            if len(file[file['Your industry'] == i]) > 0:
                numbers_industry.append(len(file[file['Your industry'] == i]))
                overal_number += len(file[file['Your industry'] == i])
        width = 0.5
        fig, ax = plt.subplots(dpi=300)
        rect = ax.bar(x, numbers_industry, width, label='Industries', color = '#0327f7')
        #ax.set_ylabel('Number of Startups')
        ax.set_xticks(x)
        ax.set_xticklabels(labels)
        for i in rect:
            height = i.get_height()
            ax.annotate('{}'.format(height),
                        xy=(i.get_x() + i.get_width() / 2, height),
                        xytext=(0, 0.2),
                        textcoords="offset points",
                        ha='center', va='bottom')
        ax.legend()
        ax.tick_params(left=False, labelleft=False, bottom=False)
        plt.tick_params(axis='x', which='major', labelsize = 10)
        fig.tight_layout()
        plt.savefig('images/industry_graph.png', bbox_inches='tight')
    
    def stages_graph(self):
        plt.clf()
        plt.close()
        global numbers_stages
        labels = ['Concept\n Stage', 'Seed\n Stage', 'Early\n Stage', 'Growth\n Stage', 'Established\n Business']
        real_labels = ['Concept Stage (got an idea)', 'Seed Stage (working on a product)', 'Early Stage (prototype ready and close to market)', 'Growth Stage (we\'re out there and making money)', 'Established Business (achieved break-even point operationally)']
        numbers_stages = []
        global stages_labels
        stages_labels = []
        global stages_numbers
        stages_numbers = []
        x = np.arange(len(labels))
        for i in real_labels:
            if len(file[file['What is the current stage of your startup?'] == i]) > 0:
                numbers_stages.append(len(file[file['What is the current stage of your startup?'] == i]))
            elif len(file[file['What is the current stage of your startup?'] == i]) == 0:
                numbers_stages.append(0)
        sorted_stages = {k:v for k,v in zip(real_labels,numbers_stages)}
        sorted_stages = sorted(sorted_stages.items(), key=lambda item: item[1])
        for k, v in sorted_stages:
            stages_labels.append(k)      
            stages_numbers.append(v)
        width = 0.7
        fig, ax = plt.subplots(dpi=300)
        rect = ax.bar(x, numbers_stages, width, label='Startups', color = '#0327f7')
        ax.set_xticks(x)
        ax.set_xticklabels(labels)
        for i in rect:
            height = i.get_height()
            ax.annotate('{}'.format(height),
                        xy=(i.get_x() + i.get_width() / 2, height),
                        xytext=(0, 0.2),
                        textcoords="offset points",
                        ha='center', va='bottom')
        ax.legend()
        ax.tick_params(left=False, labelleft=False, bottom=False, labelsize = 12)
        fig.tight_layout()
        plt.savefig('images/stages_bar.png', bbox_inches='tight')     
        
    def funding(self):
        plt.clf()
        plt.close()
        global numbers_adapted
        numbers_adapted = []
        labels = ['25k', '75k', '125k', '200k', '300k', '500k', '800k', '1M', '1.5M', '2M', '2.5M', '3M', '5M', '10M', '10+M']
        real_labels = ['1-25k', '26 - 75k', '76 - 125k', '126 - 200k', '201 - 300k', '301 - 500k', '501 - 800k', '801k - 1 M', '1 - 1.5 M', '1.5 - 2 M', '2 - 2.5 M', '2.5 - 3 M', '3 - 5 M', '5 - 10 M', '10+ M']
        x = np.arange(len(labels))
        numbers_funding = []
        global funding_labels
        funding_labels = []
        global funding_numbers
        funding_numbers = []
        for i in real_labels:
            if len(file[file['Total funding received in €'] == i]) > 0:
                numbers_funding.append(len(file[file['Total funding received in €'] == i]))
            elif len(file[file['Total funding received in €'] == i]) == 0:
                numbers_funding.append(0)
        sorted_funding = {k:v for k,v in zip(real_labels,numbers_funding)}
        sorted_funding = sorted(sorted_funding.items(), key=lambda item: item[1])
        for k, v in sorted_funding:
            funding_labels.append(k)      
            funding_numbers.append(v)      
        n = numbers_funding[0] + numbers_funding[1] + numbers_funding[2] + numbers_funding[3]
        numbers_adapted.append(n)
        n = numbers_funding[4] + numbers_funding[5] + numbers_funding[6] + numbers_funding[7]
        numbers_adapted.append(n)
        n = numbers_funding[8] + numbers_funding[9] + numbers_funding[10] + numbers_funding[11]
        numbers_adapted.append(n)
        n = numbers_funding[12] + numbers_funding[13] + numbers_funding[14]
        numbers_adapted.append(n)
        width = 0.5
        fig, ax = plt.subplots(dpi=300)
        rect = ax.bar(x, numbers_funding, width, label='Funding', color = '#0327f7')
        #ax.set_ylabel('Number of Startups')
        ax.set_xticks(x)
        ax.set_xticklabels(labels)
        for i in rect:
            height = i.get_height()
            ax.annotate('{}'.format(height),
                        xy=(i.get_x() + i.get_width() / 2, height),
                        xytext=(0, 0.2),
                        textcoords="offset points",
                        ha='center', va='bottom')
        ax.legend()
        ax.tick_params(left=False, labelleft=False, bottom=False)
        plt.tick_params(axis='x', which='major', labelsize = 9)
        fig.tight_layout()
        plt.savefig('images/funding_bar.png', bbox_inches='tight')
    
    #Doesn't Work, can't create combined graph without integers
    def combined_stages_funding(self):
        fig = plt.figure(figsize=(30,10), dpi= 80) 
        plt.title('Stages vs Funding', fontsize=22)
        sns.kdeplot(file.loc[file['What is the current stage of your startup?'] == stages_labels[-1], "Total funding received in €"], shade=True, color="deeppink", label=stages_labels[-1], alpha=.7)
        sns.kdeplot(file.loc[file['What is the current stage of your startup?'] == stages_labels[-2], "Total funding received in €"], shade=True, color="dodgerblue", label=stages_labels[-2], alpha=.7)
        sns.kdeplot(file.loc[file['What is the current stage of your startup?'] == stages_labels[-3], "Total funding received in €"], shade=True, color="g", label=stages_labels[-3], alpha=.7)
        '''x = np.arange(len(labels))
        plt.xticks(x, labels)
        plt.xlabel('Funding', fontsize=18)'''
        for i, ax in enumerate(fig.axes):
            ax.tick_params(labelleft=False, bottom = False, left = False)
        plt.savefig('images/combined_stages_funding.png', bbox_inches='tight')
        
    def product_focus_barchart(self):
        plt.clf()
        plt.close()
        global numbers_product
        labels = ['Software', 'Hardware', 'Others']
        real_labels = ['Software application', 'Physical product', 'Something else']
        x = np.arange(len(labels))
        numbers_product = []
        for i in real_labels:
            if len(file[file['Product focus'] == i]) > 0:
                numbers_product.append(len(file[file['Product focus'] == i]))
            elif len(file[file['Product focus'] == i]) == 0:
                numbers_product.append(0)
        width = 0.5
        fig, ax = plt.subplots(dpi=300)
        rect = ax.bar(x, numbers_product, width, label='Product')
        #ax.set_ylabel('Number of Startups')
        ax.set_xticks(x)
        ax.set_xticklabels(labels)
        for i in rect:
            height = i.get_height()
            ax.annotate('{}'.format(height),
                        xy=(i.get_x() + i.get_width() / 2, height),
                        xytext=(0, 0.2),
                        textcoords="offset points",
                        ha='center', va='bottom')
        ax.legend()
        ax.tick_params(left=False, labelleft=False, bottom=False)
        plt.tick_params(axis='x', which='major', labelsize = 15)
        fig.tight_layout()
        plt.savefig('images/product_focus.png', bbox_inches='tight')
        
    def product_focus_piechart(self):
        plt.clf()
        plt.close()
        global numbers_product
        labels = ['Software', 'Hardware', 'Others']
        real_labels = ['Software application', 'Physical product', 'Something else']
        numbers_product = []
        global product_labels
        product_labels = []
        global product_numbers
        product_numbers = []
        explode = (0.3, 0.17, 0.08) 
        for i in real_labels:
            if len(file[file['Product focus'] == i]) > 0:
                numbers_product.append(len(file[file['Product focus'] == i]))
            elif len(file[file['Product focus'] == i]) == 0:
                explode.pop()   
        connected = {k:v for k,v in zip(real_labels,numbers_product)}
        connected = sorted(connected.items(), key=lambda item: item[1])
        for k, v in connected:
            product_labels.append(k)
            product_numbers.append(v)
        fig1, ax1 = plt.subplots(dpi=300)
        ax1.pie(numbers_product, explode=explode, labels=labels, autopct='%1.1f%%',
                shadow=False, startangle=40)
        ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
        plt.savefig('images/product_focus_piechart.png', bbox_inches='tight')
    
    def customer_focus(self):
        plt.clf()
        plt.close()
        global numbers_customer
        labels = ['B2C', 'B2B', 'B2G']
        real_labels = ['B2C', 'B2B', 'B2G']
        x = np.arange(len(labels))
        numbers_customer = []
        for i in real_labels:
            if len(file[file['Customer focus'] == i]) > 0:
                numbers_customer.append(len(file[file['Customer focus'] == i]))
            elif len(file[file['Customer focus'] == i]) == 0:
                numbers_customer.append(0)
        width = 0.5
        fig, ax = plt.subplots(dpi=300)
        rect = ax.bar(x, numbers_customer, width, label='Customers')
        #ax.set_ylabel('Number of Startups')
        ax.set_xticks(x)
        ax.set_xticklabels(labels)
        for i in rect:
            height = i.get_height()
            ax.annotate('{}'.format(height),
                        xy=(i.get_x() + i.get_width() / 2, height),
                        xytext=(0, 0.2),
                        textcoords="offset points",
                        ha='center', va='bottom')
        ax.legend()
        ax.tick_params(left=False, labelleft=False, bottom=False)
        plt.tick_params(axis='x', which='major', labelsize = 15)
        fig.tight_layout()
        plt.savefig('images/customer_focus.png', bbox_inches='tight')
        
    def customer_focus_piechart(self):
        plt.clf()
        plt.close()
        global numbers_customer
        labels = ['B2C', 'B2B', 'B2G']
        real_labels = []
        numbers_product = []
        global customer_labels
        customer_labels = []
        global customer_numbers
        customer_numbers= []
        explode = [0.3, 0.17, 0.08]
        for i in labels:
            if len(file[file['Customer focus'] == i]) > 0:
                numbers_product.append(len(file[file['Customer focus'] == i]))
                real_labels.append(i)
            elif len(file[file['Customer focus'] == i]) == 0:
                explode.pop()
        connected = {k:v for k,v in zip(real_labels,numbers_product)}
        connected = sorted(connected.items(), key=lambda item: item[1])
        for k, v in connected:
            customer_labels.append(k)
            customer_numbers.append(v)
        fig1, ax1 = plt.subplots(dpi=300)
        ax1.pie(numbers_product, explode=explode, labels=real_labels, autopct='%1.1f%%',
                shadow=False, startangle=40)
        ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
        plt.savefig('images/customer_focus_piechart.png', bbox_inches='tight')

    def country_graph(self):
        plt.clf()
        plt.close()
        real_labels = []
        global ordered_labels
        ordered_labels = []
        numbers = []  
        global ordered_numbers
        ordered_numbers = []
        group = file.groupby('Country of incorporation / registration').size()
        for i, v in group.items():
            if ',' not in i:
                if i not in real_labels:
                    real_labels.append(i)
                    numbers.append(v)
            if ',' in  i:
                new = i.split(',', 1)[0].replace(',', '')
                if new not in real_labels:
                    real_labels.append(new)
                    numbers.append(v)
        connected = {k:v for k,v in zip(real_labels,numbers)}
        connected = sorted(connected.items(), key=lambda item: item[1])
        for k, v in connected:
            ordered_labels.append(k)
            ordered_numbers.append(v)
        global map_data
        map_data = pd.DataFrame(data={"country": ordered_labels, "number": ordered_numbers})
        x = np.arange(len(ordered_labels))
        width = 0.5
        fig, ax = plt.subplots(figsize=(7,3), dpi=300)
        rect = ax.bar(x, ordered_numbers, width, color = '#0327f7')
        ax.set_xticks(x)
        ax.set_xticklabels(ordered_labels, rotation=45, horizontalalignment='right')
        for i in rect:
            height = i.get_height()
            ax.annotate('{}'.format(height),
                        xy=(i.get_x() + i.get_width() / 2, height),
                        xytext=(0, 0),
                        textcoords="offset points",
                        ha='center', va='bottom')
        ax.tick_params(left=False, labelleft=False, bottom=False)
        '''ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_visible(False)'''
        plt.tick_params(axis='x', which='major', labelsize = 9)
        fig.tight_layout()
        plt.savefig('images/country_graph.png', bbox_inches='tight')
        plt.clf()
        plt.close()
        
    def country_map(self):
        shapefile = 'map/ne_110m_admin_0_countries.shp'
        gdf = gpd.read_file(shapefile)[['ADMIN', 'ADM0_A3', 'geometry']]
        gdf.columns = ['country', 'country_code', 'geometry'] 
        gdf = gdf.drop(gdf.index[159]) #drop Antarctica
        merged = gdf.merge(map_data, left_on = 'country', right_on = 'country', how = 'left')
        merged.fillna('No data', inplace = True)
        merged_json = json.loads(merged.to_json())
        json_data = json.dumps(merged_json)
        geosource = GeoJSONDataSource(geojson = json_data)
        palette = brewer['Blues'][8]
        palette = palette[::-1]
        color_mapper = LinearColorMapper(palette = palette, low = 0, high = ordered_numbers[-1], nan_color = '#EAEAEA')
        color_bar = ColorBar(color_mapper=color_mapper, label_standoff=8,width = 500, height = 20,
        border_line_color=None,location = (0,0), orientation = 'horizontal')
        p = figure(title = 'Share of startups in countries', plot_height = 600 , plot_width = 950, toolbar_location = None)
        p.xgrid.grid_line_color = None
        p.ygrid.grid_line_color = None
        p.axis.visible = False
        p.patches('xs','ys', source = geosource,fill_color = {'field' :'number', 'transform' : color_mapper},
                  line_color = 'black', line_width = 0.25, fill_alpha = 1)
        p.add_layout(color_bar, 'below')
        export_png(p, filename="map.png")
        
    def word_cloud(self):
        plt.clf()
        plt.close()
        wordcloud_mask = np.array(Image.open('images/pionerd.png'))
        wordcloud_mask[wordcloud_mask == 0] = 255
        stopwords = set(STOPWORDS)
        stopwords.update(['client', 'customer', 'use', 'service', 'based', 'product', 'will', 'is', 'are', 'offer', 'company', 'use', 'project', 'provide', 'new'])
        wordcloud = WordCloud(mask = wordcloud_mask, stopwords=stopwords, max_words=80, background_color="white", contour_width=5, contour_color='black').generate(' '.join(file['Description']))
        plt.figure(figsize=[8,8])
        plt.imshow(wordcloud, interpolation='bilinear')
        plt.axis("off")
        wordcloud.to_file("images/wordcloud.png")
        
# Step 3 
# Using docx python library to create Word file and input all information.
def first_page():
    global document
    document = Document()
    section = document.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    r= p.add_run()
    r.add_picture('images/pioneers_logo.png', width=Inches(0.9), height=Inches(0.5))
    p.alignment = 1
    for i in range(10):
        document.add_paragraph()
    p = document.add_paragraph()
    r = p.add_run()
    r.add_text('Open Call Report')
    r.font.size = Pt(45)
    r.bold = True
    p = document.add_paragraph()
    r = p.add_run()
    r.add_text('For ' + open_call)
    r.font.size = Pt(25)
    '''r = p.add_run()
    r.add_picture('images/wordcloud.png', width=Inches(6), height=Inches(6))'''
    r.add_break(WD_BREAK.PAGE)

def ecosystem_map():   
    document.add_paragraph('Industries and Operations', style='Title')
    p = document.add_paragraph()
    p.add_run('Successful collaboration begins with choosing the right startups for your innovation initiatives! The top 4 industries of ' + open_call + ' open call are:')
    for i in industries:
        document.add_paragraph(i, style='List Bullet')
    p = document.add_paragraph()
    r = p.add_run()
    r.add_picture('images/ecosystem_map.png', width=Inches(5), height=Inches(3.2))
    p.alignment = 1

def industry():
    '''r.add_text('Areas of Operations')
    r.bold = True
    r.font.size = Pt(16)
    p.alignment = 1'''
    p = document.add_paragraph()
    r = p.add_run()
    table = document.add_table(1, 2)
    cells = table.rows[0].cells
    paragraph = cells[0].paragraphs[0]
    run = paragraph.add_run()
    run.add_text('There is a total number of {} startups from top {} industries:'.format(overal_number, len(industries), numbers_industry[0], industries[0], numbers_industry[1], industries[1], numbers_industry[2], industries[2], numbers_industry[3], industries[3]))
    for i in range(len(industries)):
        cells[0].add_paragraph('• {} from {}'.format(numbers_industry[i], industries[i]))
    paragraph = cells[1].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture('images/industry_graph.png', width = Inches(2.95), height = Inches(2.2))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)

def country():
    document.add_paragraph('Countries of Operations', style='Title')
    p = document.add_paragraph()
    r = p.add_run()
    r.add_picture('images/country_graph.png', width = Inches(5.5), height = Inches(2.4))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph()
    r = p.add_run()
    r.add_text('Top 3 countries: {}, {} & {}'.format(ordered_labels[-1],ordered_labels[-2],ordered_labels[-3]))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()
    '''p = document.add_paragraph()
    r = p.add_run()
    r.add_text('World Map of Startups')
    r.bold = True
    r.font.size = Pt(16)
    p.alignment = 1
    document.add_paragraph()'''
    p = document.add_paragraph()
    r = p.add_run()
    r.add_picture('map.png', width=Inches(6.2), height=Inches(3.7))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r.add_break(WD_BREAK.PAGE)
    
def stages():
    p = document.add_paragraph()
    r = p.add_run()
    r.add_text('Stages - from Idea to Established Business')
    r.bold = True
    r.font.size = Pt(16)
    table = document.add_table(1, 2)
    cells = table.rows[0].cells
    paragraph = cells[0].paragraphs[0]
    run = paragraph.add_run()
    run.add_text('• {} Concept Stage'.format(numbers_stages[0]))
    cells[0].add_paragraph('• {} Seed Stage'.format(numbers_stages[1]))
    cells[0].add_paragraph('• {} Early Stage'.format(numbers_stages[2]))
    cells[0].add_paragraph('• {} Growth Stage'.format(numbers_stages[3]))
    cells[0].add_paragraph('• {} Established Businesses'.format(numbers_stages[4]))
    paragraph = cells[1].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture('images/stages_bar.png', width = Inches(2.95), height = Inches(2))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
def funding():
    p = document.add_paragraph()
    r = p.add_run()
    r.add_text('Funding')
    r.bold = True
    r.font.size = Pt(16)
    table = document.add_table(1, 2)
    cells = table.rows[0].cells
    paragraph = cells[0].paragraphs[0]
    run = paragraph.add_run()
    cells[0].add_paragraph
    run.add_text('• {} Pre-Seed Startups (1-200k)'.format(numbers_adapted[0]))
    cells[0].add_paragraph('• {} Seed Startups (200- 1 Mil)'.format(numbers_adapted[1]))
    cells[0].add_paragraph('• {} Series A Startups (1-3  Mil)'.format(numbers_adapted[2]))
    cells[0].add_paragraph('• {} Series B Startups (>3 Mil)'.format(numbers_adapted[3]))
    paragraph = cells[1].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture('images/funding_bar.png', width = Inches(2.95), height = Inches(2))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def stages_funding():
    document.add_paragraph('Facts and Figures', style='Title')
    p = document.add_paragraph()
    r = p.add_run()
    r.add_text('Stages & Funding')
    r.bold = True
    r.font.size = Pt(16)
    table = document.add_table(2, 2)
    cells = table.rows[0].cells
    run = cells[0].paragraphs[0].add_run('Stages')
    run.bold = True
    run.alingment = 1
    run = cells[1].paragraphs[0].add_run('Funding')
    run.bold = True
    run.alingment = 1
    cells = table.rows[1].cells
    paragraph = cells[0].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture('images/stages_bar.png', width = Inches(2.95), height = Inches(2))
    paragraph = cells[1].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture('images/funding_bar.png', width = Inches(2.95), height = Inches(2))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()
    p = document.add_paragraph()
    r = p.add_run()
    r.add_text('There are {} {} startups in the open call. The most common funding is between {} and {} startups received it.'.format(stages_numbers[-1], stages_labels[-1], funding_labels[-1], funding_numbers[-1]))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
def combined_s_f():
    p = document.add_paragraph()
    r = p.add_run()
    r.add_text('Combined')
    r.bold = True
    r.font.size = Pt(16)
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph()
    r = p.add_run()
    r.add_picture('images/combined_stages_funding.png', width = Inches(6), height = Inches(2))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
# This version has descriptions of products and customers
def product():
    p = document.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)
    document.add_paragraph('Product and Customer Focus', style='Title')
    p = document.add_paragraph()
    r = p.add_run()
    r.add_text('Product Focus')
    r.bold = True
    r.font.size = Pt(16)
    table = document.add_table(1, 2)
    cells = table.rows[0].cells
    cells[0].add_paragraph('• {} Software Startups'.format(numbers_product[0]))
    cells[0].add_paragraph('• {} Hardware Startups'.format(numbers_product[1]))
    cells[0].add_paragraph('• {} Other Startups'.format(numbers_product[2]))
    paragraph = cells[1].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture('images/product_focus_piechart.png', width = Inches(2.95), height = Inches(2))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def customer():
    p = document.add_paragraph()
    r = p.add_run()
    r.add_text('Customer Focus')
    r.bold = True
    r.font.size = Pt(16)
    table = document.add_table(1, 2)
    cells = table.rows[0].cells
    cells[0].add_paragraph('• {} B2C'.format(numbers_customer[0]))
    cells[0].add_paragraph('• {} B2B'.format(numbers_customer[1]))
    cells[0].add_paragraph('• {} B2G'.format(numbers_customer[2]))
    paragraph = cells[1].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture('images/customer_focus_piechart.png', width = Inches(2.95), height = Inches(2))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)
    
# This includes product and customers in one table
def product_customer():
    document.add_paragraph()
    p = document.add_paragraph()
    r = p.add_run()
    r.add_text('Products & Customers')
    r.bold = True
    r.font.size = Pt(16)
    table = document.add_table(2, 2)
    cells = table.rows[0].cells
    run = cells[0].paragraphs[0].add_run('Product Focus')
    run.bold = True
    run.alingment = 1
    run = cells[1].paragraphs[0].add_run('Customer Focus')
    run.bold = True
    run.alingment = 1
    cells = table.rows[1].cells
    paragraph = cells[0].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture('images/product_focus_piechart.png', width = Inches(2.95), height = Inches(2))
    paragraph = cells[1].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture('images/customer_focus_piechart.png', width = Inches(2.95), height = Inches(2))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()
    p = document.add_paragraph()
    r = p.add_run()
    product_sum = 0
    for i in numbers_product:
        product_sum += i
    r.add_text('{} from {} applicants have {} product. {} startups focus on {} customers.'.format(product_numbers[-1],product_sum,product_labels[-1],customer_numbers[-1],customer_labels[-1]))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)
    
def set_column_width(column, width):
    column.width = width
    for cell in column.cells:
        cell.width = width

def table_maker(_):
    p = document.add_paragraph()
    p.alignment = 1
    r = p.add_run()
    r.font.size = Pt(16)
    r.bold = True
    r.add_text(industries[_])
    table = document.add_table(1, 3)
    heading_cells = table.rows[0].cells
    heading_cells[0].text = 'Company Name  '
    heading_cells[1].text = 'Domain  '
    heading_cells[2].text = 'Description'
    new = file.loc[file['Your industry'] == industries[_]]
    keep = ['Company Name', 'Domain', 'Description']
    for name, domain, descr in new[keep].itertuples(index=False): 
        cells = table.add_row().cells
        if len(name) > 16:
            name = name[0:16] + '..'
            cells[0].text = name
        if len(name) <= 16:
            cells[0].text = name
        cells[1].text = domain  
        if len(descr) > 60:
            descr = descr[0:60] + '...'
            cells[2].text = descr
        if len(descr) <= 60:
            cells[2].text = descr   
    set_column_width(table.columns[0], Inches(2.5))
    set_column_width(table.columns[1], Inches(2.5))
    set_column_width(table.columns[2], Inches(5))
    table.style = 'Table Grid'
    document.add_paragraph()

def launch():
    global open_call
    print('Write the name of the open call you downloaded:')
    open_call = input()
    print('Executing the script, wait 15 seconds.')
    plt.style.use('seaborn-darkgrid')
    csv_from_excel()
    remove_cols()
    global file
    file = pd.read_csv('csv_files/new_csv.csv', index_col=0)    
    file['Your industry'] = file['Your industry'].apply(coma_remove)
    global industries
    industries = []
    industries_all = []
    size = []
    every = []
    for i in file['Your industry']:
        if i not in every:
            every.append(i)
            industries_all.append(i)
            size.append(len(file[file['Your industry'] == i]))
    industries_all = {k:v for k,v in zip(industries_all,size)}
    industries_all = sorted(industries_all.items(), key=lambda item: item[1], reverse = True)
    for k, v in industries_all:
        if len(industries) == 4:
            break
        industries.append(k)
    g = graph()
    g.boxes()
    g.industry_graph()
    g.stages_graph()
    g.funding()
    g.product_focus_piechart()
    g.customer_focus_piechart()
    g.country_graph()
    g.country_map()
    #g.combined_stages_funding()
    #g.word_cloud()
    first_page()
    ecosystem_map()
    industry()
    country()
    #stages()
    #funding()
    stages_funding()
    #combined_s_f()
    #product()
    #customer()
    product_customer()
    document.add_paragraph('Startup Export', style='Title')
    for i in range(0, 3):
        table_maker(i)
    document.save('Report.docx')
    print('The script is successfully completed. Denis is happy!')

# Need to figure out how to change .docx into .pdf
if __name__ == '__main__':    
    launch()